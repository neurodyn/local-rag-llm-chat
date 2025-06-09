import os
import logging
from datetime import datetime
from typing import List, Optional, Dict, Tuple, Any
import uuid
from pathlib import Path
import magic
from docx import Document as DocxDocument
from PyPDF2 import PdfReader
from huggingface_hub import hf_hub_download, HfApi
from tqdm.auto import tqdm
from .mlx_pipeline import MLXPipeline
from .chat_mlx import ChatMLX
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_core.messages import HumanMessage, AIMessage
from langchain_community.tools import DuckDuckGoSearchRun
from fastapi import UploadFile, HTTPException
from starlette.concurrency import run_in_threadpool
from ..models.document import Document
from unstructured.partition.pdf import partition_pdf
import asyncio
from duckduckgo_search import DDGS
from serpapi import GoogleSearch
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class WebSearch:
    def __init__(self, search_engine: str = "duckduckgo"):
        """
        Initialize web search with specified engine.
        
        Args:
            search_engine (str): Either "duckduckgo" or "serpapi"
        """
        self.search_engine = search_engine.lower()
        logger.info(f"Initializing WebSearch with engine: {self.search_engine}")
        
        # Initialize search engine
        if self.search_engine == "duckduckgo":
            self.ddgs = DDGS()
            logger.info("DuckDuckGo search engine initialized")
        elif self.search_engine == "serpapi":
            self.serpapi_key = os.getenv("SERPAPI_API_KEY")
            if not self.serpapi_key:
                logger.warning("SERPAPI_API_KEY not found in environment variables. Falling back to DuckDuckGo.")
                self.search_engine = "duckduckgo"
                self.ddgs = DDGS()
                logger.info("Falling back to DuckDuckGo due to missing SERPAPI_API_KEY")
            else:
                logger.info("SerpAPI search engine initialized with API key")
        else:
            logger.warning(f"Invalid search engine '{search_engine}'. Defaulting to DuckDuckGo.")
            self.search_engine = "duckduckgo"
            self.ddgs = DDGS()
            logger.info("Defaulting to DuckDuckGo due to invalid search engine")
    
    def run(self, query: str) -> str:
        """
        Run a web search query using the configured search engine.
        
        Args:
            query (str): The search query
            
        Returns:
            str: Formatted search results or empty string on error
        """
        try:
            logger.info(f"Running search with engine: {self.search_engine}")
            if self.search_engine == "duckduckgo":
                return self._search_duckduckgo(query)
            else:
                return self._search_serpapi(query)
        except Exception as e:
            logger.error(f"Web search error: {e}")
            return ""
    
    def _search_duckduckgo(self, query: str) -> str:
        """Search using DuckDuckGo."""
        try:
            logger.info(f"Executing DuckDuckGo search for query: {query}")
            results = list(self.ddgs.text(query, max_results=3))
            logger.info(f"DuckDuckGo raw search results: {results}")
            
            formatted_results = []
            for result in results:
                formatted_results.append(f"{result['title']}\n{result['body']}")
            
            formatted_output = "\n\n".join(formatted_results) if formatted_results else "No results found."
            logger.info(f"DuckDuckGo formatted output:\n{formatted_output}")
            return formatted_output
        except Exception as e:
            logger.error(f"DuckDuckGo search error: {e}")
            return ""
    
    def _search_serpapi(self, query: str) -> str:
        """Search using SerpAPI."""
        try:
            if not self.serpapi_key:
                logger.warning("No SERPAPI_API_KEY found. Falling back to DuckDuckGo.")
                return self._search_duckduckgo(query)
            
            logger.info(f"Executing SerpAPI search for query: {query}")
            search = GoogleSearch({
                "q": query,
                "api_key": self.serpapi_key,
                "num": 3  # Number of results
            })
            results = search.get_dict()
            logger.info(f"SerpAPI raw search results: {results}")
            
            formatted_results = []
            if "organic_results" in results:
                for result in results["organic_results"][:3]:
                    title = result.get("title", "")
                    snippet = result.get("snippet", "")
                    formatted_results.append(f"{title}\n{snippet}")
            
            formatted_output = "\n\n".join(formatted_results) if formatted_results else "No results found."
            logger.info(f"SerpAPI formatted output:\n{formatted_output}")
            return formatted_output
        except Exception as e:
            logger.error(f"SerpAPI search error: {e}")
            # Fall back to DuckDuckGo on error
            logger.info("Falling back to DuckDuckGo search due to SerpAPI error")
            return self._search_duckduckgo(query)

class RAGService:
    def __init__(self, search_engine: str = "duckduckgo"):
        """
        Initialize the RAG service with basic setup.
        
        Args:
            search_engine (str): Either "duckduckgo" or "serpapi" for web search
        """
        self.initialization_error = None
        self.initialization_status = "Not started"
        self.initialization_progress = {
            "current_step": "",
            "total_steps": 4,
            "current_step_number": 0,
            "download_progress": 0,
            "detailed_status": ""
        }
        self.documents_dir = Path("data/documents")
        self.vector_store_dir = Path("data/vectorstore")
        self.init_state_file = Path("data/init_state.json")
        self.documents_dir.mkdir(parents=True, exist_ok=True)
        self.vector_store_dir.mkdir(parents=True, exist_ok=True)
        
        # Load initialization state
        self.is_initialized = self._load_init_state()
        
        self.llm = None
        self.chat_model = None
        self.embeddings = None
        self.vector_store = None
        self.conversations = {}
        self.search_tool = None
        self.search_engine = search_engine.lower()
        logger.info(f"RAGService initialized with search engine: {self.search_engine}")
        
        # Supported file types
        self.supported_mimetypes = {
            'application/pdf': self._extract_pdf_text,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._extract_docx_text,
            'text/plain': self._extract_text_file,
        }

    def _load_init_state(self) -> bool:
        """Load initialization state from file."""
        try:
            if self.init_state_file.exists():
                import json
                with open(self.init_state_file, 'r') as f:
                    state = json.load(f)
                    return state.get('initialized', False)
            return False
        except Exception as e:
            logger.error(f"Error loading initialization state: {e}")
            return False

    def _save_init_state(self, initialized: bool):
        """Save initialization state to file."""
        try:
            import json
            with open(self.init_state_file, 'w') as f:
                json.dump({'initialized': initialized}, f)
        except Exception as e:
            logger.error(f"Error saving initialization state: {e}")

    async def _load_models(self):
        """Load models if already initialized."""
        try:
            model_id = "mlx-community/Mistral-7B-Instruct-v0.3-4bit"
            embedding_model_id = "all-MiniLM-L6-v2"
            
            # Load MLX model
            self.llm = await run_in_threadpool(
                MLXPipeline.from_model_id,
                model_id,
                pipeline_kwargs={"max_tokens": 512, "temperature": 0.7}
            )
            self.chat_model = ChatMLX(llm=self.llm)
            
            # Load embeddings
            self.embeddings = await run_in_threadpool(
                lambda: HuggingFaceEmbeddings(
                    model_name=embedding_model_id,
                    model_kwargs={'device': 'cpu'}
                )
            )
            
            # Load vector store
            self.vector_store = await run_in_threadpool(
                lambda: Chroma(
                    persist_directory=str(self.vector_store_dir),
                    embedding_function=self.embeddings
                )
            )
            
            logger.info("Models loaded successfully")
        except Exception as e:
            logger.error(f"Error loading models: {e}")
            self.is_initialized = False
            self._save_init_state(False)
            raise

    def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from a PDF file with enhanced table support."""
        try:
            logger.info(f"Extracting text from PDF: {file_path}")
            elements = partition_pdf(str(file_path))
            text = ""
            
            for element in elements:
                if hasattr(element, "text"):
                    text += element.text + "\n\n"
                    
            if not text.strip():
                raise ValueError("No text could be extracted from the PDF")
                
            logger.info(f"Successfully extracted {len(text)} characters from PDF")
            return text
                
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {e}")
            raise HTTPException(
                status_code=500,
                detail=f"Error processing PDF file: {str(e)}"
            )

    def _extract_docx_text(self, file_path: Path) -> str:
        """Extract text from a Word document with enhanced table support."""
        try:
            doc = DocxDocument(file_path)
            text = ""
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
                
            # Extract tables
            for table in doc.tables:
                text += "\nTable:\n"
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    # Only add rows that have actual content
                    if any(cell for cell in row_text):
                        text += " | ".join(row_text) + "\n"
                text += "\n"
                
            return text
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {e}")
            raise HTTPException(
                status_code=500,
                detail=f"Error processing Word document: {str(e)}"
            )

    def _extract_text_file(self, file_path: Path) -> str:
        """Extract text from a plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try with different encodings if UTF-8 fails
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    return f.read()
            except Exception as e:
                logger.error(f"Error reading text file {file_path}: {e}")
                raise HTTPException(
                    status_code=500,
                    detail=f"Error processing text file: {str(e)}"
                )

    async def update_status(self, step: str, progress: int = 0, details: str = ""):
        """Update initialization status with progress information."""
        steps = {
            "Checking model files": 0,
            "Downloading Mistral model": 1,
            "Setting up embeddings": 2,
            "Configuring vector store": 3,
            "Initialization complete": 4
        }
        
        # Run the state update in a threadpool to ensure thread safety
        await run_in_threadpool(self._update_status_sync, step, steps.get(step, 0), progress, details)
        
    def _update_status_sync(self, step: str, step_number: int, progress: int, details: str):
        """Synchronous helper for updating status in a thread-safe way."""
        self.initialization_progress["current_step"] = step
        self.initialization_progress["current_step_number"] = step_number
        self.initialization_progress["download_progress"] = progress
        self.initialization_progress["detailed_status"] = details
        self.initialization_status = f"{step} - {details}"
        logger.info(self.initialization_status)

    async def initialize(self):
        """Asynchronously initialize the MLX model and vector store."""
        try:
            if self.is_initialized:
                return

            # Step 1: Check model cache and download if needed
            await self.update_status(
                "Checking model files",
                10,
                "Verifying model cache and downloading if needed..."
            )

            model_id = "mlx-community/Mistral-7B-Instruct-v0.3-4bit"
            embedding_model_id = "all-MiniLM-L6-v2"

            # Get model file sizes for progress tracking
            api = HfApi()
            try:
                await self.update_status(
                    "Checking model files",
                    20,
                    "Checking model repository information..."
                )
                model_files = await run_in_threadpool(api.list_repo_files, model_id)
                total_size = await run_in_threadpool(
                    lambda: sum(
                        api.get_repo_info(model_id).siblings[i].size 
                        for i, file in enumerate(model_files) 
                        if file.endswith('.bin') or file.endswith('.json')
                    )
                )
            except Exception as e:
                logger.warning(f"Could not get model file sizes: {e}")
                total_size = None

            # Download and initialize MLX model
            await self.update_status(
                "Downloading Mistral model",
                30,
                "Starting model download..."
            )
            
            try:
                # Initialize MLX model through LangChain
                await self.update_status(
                    "Downloading Mistral model",
                    40,
                    "Downloading model files..."
                )
                
                self.llm = await run_in_threadpool(
                    MLXPipeline.from_model_id,
                    model_id,
                    pipeline_kwargs={"max_tokens": 512, "temperature": 0.7}
                )
                
                await self.update_status(
                    "Downloading Mistral model",
                    60,
                    "Model downloaded, setting up MLX pipeline..."
                )
                
                # Create ChatMLX model for better prompt handling
                self.chat_model = ChatMLX(llm=self.llm)
                
            except Exception as e:
                await self.update_status(
                    "Error downloading model",
                    0,
                    f"Failed to download model: {str(e)}"
                )
                raise

            # Initialize embeddings
            await self.update_status(
                "Setting up embeddings",
                70,
                "Downloading and initializing embedding model..."
            )
            
            self.embeddings = await run_in_threadpool(
                lambda: HuggingFaceEmbeddings(
                    model_name=embedding_model_id,
                    model_kwargs={'device': 'cpu'}
                )
            )
            
            await self.update_status(
                "Setting up embeddings",
                80,
                "Embedding model initialized successfully"
            )
            
            # Initialize vector store
            await self.update_status(
                "Configuring vector store",
                90,
                "Setting up document storage..."
            )
            
            # Initialize or load existing vector store
            self.vector_store = await run_in_threadpool(
                lambda: Chroma(
                    persist_directory=str(self.vector_store_dir),
                    embedding_function=self.embeddings,
                    collection_metadata={"hnsw:space": "cosine"}  # Use cosine similarity
                )
            )
            
            await self.update_status(
                "Initialization complete",
                100,
                "System is ready to use"
            )
            
            self.is_initialized = True
            self._save_init_state(True)
            
        except Exception as e:
            self.initialization_error = str(e)
            self.initialization_status = f"Initialization failed: {str(e)}"
            self.initialization_progress["detailed_status"] = str(e)
            logger.error(self.initialization_status)
            self.is_initialized = False
            self._save_init_state(False)
            raise

    def get_initialization_status(self) -> Dict[str, any]:
        """Get the current initialization status with detailed progress."""
        return {
            "is_initialized": self.is_initialized,
            "status": self.initialization_status,
            "error": self.initialization_error,
            "progress": self.initialization_progress
        }

    async def ensure_initialized(self):
        """Ensure the service is initialized."""
        if not self.is_initialized:
            await self.initialize()
        elif self.llm is None or self.chat_model is None or self.embeddings is None or self.vector_store is None:
            await self._load_models()
        
        # Initialize search tool if not already done
        if self.search_tool is None:
            logger.info(f"Initializing new search tool with engine: {self.search_engine}")
            self.search_tool = WebSearch(search_engine=self.search_engine)
            logger.info(f"Search tool initialized with engine: {self.search_tool.search_engine}")

    def list_documents(self) -> List[Document]:
        """List all processed documents."""
        try:
            documents = []
            for file_path in self.documents_dir.glob("*"):
                if file_path.is_file():
                    doc_id = file_path.name.split("_")[0]
                    filename = "_".join(file_path.name.split("_")[1:])
                    documents.append(
                        Document(
                            id=doc_id,
                            filename=filename,
                            upload_time=datetime.fromtimestamp(file_path.stat().st_mtime),
                            content_type="application/octet-stream",
                            vector_store_id=doc_id
                        )
                    )
            return documents
        except Exception as e:
            logger.error(f"Error listing documents: {str(e)}")
            raise HTTPException(status_code=500, detail="Error listing documents")

    async def process_document(self, file: UploadFile) -> Document:
        """Process an uploaded document and add it to the vector store."""
        await self.ensure_initialized()
        try:
            logger.info(f"Processing document: {file.filename}")
            
            # Generate unique ID for the document
            doc_id = str(uuid.uuid4())
            
            # Save the file
            file_path = self.documents_dir / f"{doc_id}_{file.filename}"
            content = await file.read()
            file_path.write_bytes(content)
            
            # Detect file type
            mime_type = magic.from_file(str(file_path), mime=True)
            logger.info(f"Detected MIME type: {mime_type}")
            
            if mime_type not in self.supported_mimetypes:
                raise HTTPException(
                    status_code=400,
                    detail=f"Unsupported file type: {mime_type}. Supported types are PDF, Word (docx), and text files."
                )
            
            # Extract text using appropriate method
            text = self.supported_mimetypes[mime_type](file_path)
            if not text.strip():
                raise ValueError("No text could be extracted from the document")
            
            logger.info(f"Extracted {len(text)} characters of text")
            
            # Determine optimal chunk size and overlap based on document type and content
            chunk_size = 1000  # Default
            chunk_overlap = 200  # Default
            
            # Adjust based on document type and content length
            if mime_type == 'application/pdf':
                # PDFs often have natural page breaks, use larger chunks
                chunk_size = 1500
                chunk_overlap = 300
            elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                # Word docs might have more structured content
                chunk_size = 1200
                chunk_overlap = 250
            elif len(text) < 5000:
                # For small documents, use smaller chunks
                chunk_size = 500
                chunk_overlap = 100
            
            # Create text splitter with custom settings
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                length_function=len,
                separators=["\n\n", "\n", " ", ""]  # Try natural breaks first
            )
            
            # Split text into chunks
            texts = text_splitter.split_text(text)
            logger.info(f"Split text into {len(texts)} chunks (size: {chunk_size}, overlap: {chunk_overlap})")
            
            # Add to vector store with improved metadata
            try:
                self.vector_store.add_texts(
                    texts,
                    metadatas=[{
                        "source": str(file_path),
                        "doc_id": doc_id,
                        "chunk_id": i,
                        "total_chunks": len(texts),
                        "chunk_size": chunk_size,
                        "chunk_overlap": chunk_overlap,
                        "mime_type": mime_type,
                        "filename": file.filename
                    } for i in range(len(texts))],
                    ids=[f"{doc_id}_{i}" for i in range(len(texts))]
                )
                
                # Verify document was indexed with MMR search
                retriever = self.vector_store.as_retriever(
                    search_type="mmr",
                    search_kwargs={
                        "filter": {"doc_id": doc_id},
                        "k": 2,
                        "fetch_k": 10,
                        "lambda_mult": 0.7
                    }
                )
                verification_docs = retriever.get_relevant_documents("test")
                if not verification_docs:
                    raise ValueError("Document was not properly indexed in vector store")
                    
                logger.info(f"Successfully indexed {len(texts)} chunks in vector store")
                
            except Exception as e:
                logger.error(f"Error adding document to vector store: {e}")
                # Clean up the saved file
                file_path.unlink(missing_ok=True)
                raise HTTPException(
                    status_code=500,
                    detail=f"Error indexing document: {str(e)}"
                )
            
            # Initialize conversation history
            self.conversations[doc_id] = []
            
            # Create and return document metadata
            document = Document(
                id=doc_id,
                filename=file.filename,
                upload_time=datetime.now(),
                content_type=mime_type,
                vector_store_id=doc_id
            )
            
            logger.info(f"Document processed successfully: {doc_id}")
            return document
            
        except HTTPException as e:
            raise e
        except Exception as e:
            logger.error(f"Error processing document {file.filename}: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Error processing document: {str(e)}")

    def _extract_response_text(self, response) -> str:
        """Helper method to extract text from model response."""
        if hasattr(response, 'generations'):
            # Handle ChatResult
            if len(response.generations) > 0:
                generation = response.generations[0]
                if hasattr(generation, 'message'):
                    return generation.message.content
                elif hasattr(generation, 'text'):
                    return generation.text
        elif isinstance(response, dict) and 'content' in response:
            return response['content']
        elif isinstance(response, (list, tuple)) and len(response) > 0:
            return str(response[0])
        elif hasattr(response, 'content'):
            return response.content
        elif hasattr(response, 'text'):
            return response.text
        else:
            return str(response)

    async def process_query_multi(self, document_ids: List[str], query: str) -> Dict[str, Any]:
        """Process a query across multiple documents using RAG."""
        await self.ensure_initialized()
        try:
            logger.info(f"Processing query across documents {document_ids}: {query}")
            
            # Check if this is a web search request
            should_web_search = any(keyword in query.lower() for keyword in [
                "search online", "search the internet","search the web", "search internet", "look up online",
                "find online", "search for", "look for"
            ])
            
            context_parts = []
            sources = []
            
            # If web search is requested, perform the search with retry logic
            if should_web_search:
                try:
                    logger.info("Performing web search")
                    # Add exponential backoff for rate limits
                    max_retries = 3
                    base_delay = 2  # seconds
                    
                    for attempt in range(max_retries):
                        try:
                            # Add delay that increases with each retry
                            if attempt > 0:
                                delay = base_delay * (2 ** (attempt - 1))
                                logger.info(f"Retrying web search after {delay} seconds (attempt {attempt + 1}/{max_retries})")
                                await asyncio.sleep(delay)
                            
                            search_result = await run_in_threadpool(self.search_tool.run, query)
                            if search_result:
                                context_parts.append(f"Web Search Results:\n{search_result}")
                                sources.append({
                                    "content": search_result,
                                    "metadata": {"source": "web_search", "engine": self.search_engine}
                                })
                                logger.info("Web search completed successfully")
                                break
                        except Exception as search_error:
                            if "Ratelimit" in str(search_error) and attempt < max_retries - 1:
                                logger.warning(f"Web search rate limited (attempt {attempt + 1}/{max_retries})")
                                continue
                            else:
                                logger.warning(f"Web search failed: {search_error}")
                                context_parts.append("Note: Web search was unavailable due to rate limiting. Proceeding with available document context only.")
                                break
                except Exception as e:
                    logger.error(f"Web search error: {e}")
                    # Continue with document search even if web search fails
            
            # Get relevant documents from the vector store if document_ids are provided
            if document_ids and self.vector_store is not None:
                retriever = self.vector_store.as_retriever(
                    search_type="mmr",
                    search_kwargs={
                        "filter": {"doc_id": {"$in": document_ids}},
                        "k": 6,
                        "fetch_k": 20,
                        "lambda_mult": 0.7
                    }
                )
                
                relevant_docs = retriever.get_relevant_documents(query)
                if relevant_docs:
                    context_parts.append("Document Context:\n" + "\n\n".join([doc.page_content for doc in relevant_docs]))
                    
                    # Add document sources
                    for doc in relevant_docs:
                        sources.append({
                            "content": doc.page_content,
                            "metadata": doc.metadata
                        })
            
            # Combine all context parts
            context = "\n\n".join(context_parts) if context_parts else "No context available."
            
            # Create prompt
            prompt = f"""Use the following information to answer the question. If you cannot find the answer in the provided information, say "I don't have enough information to answer that."

Information:
{context}

Question: {query}

Answer:"""
            
            # Get response from model
            response = self.chat_model.invoke(prompt)
            response_text = self._extract_response_text(response)
            response_text = response_text.strip()
            
            # Update conversation history for all involved documents
            for doc_id in document_ids:
                if doc_id not in self.conversations:
                    self.conversations[doc_id] = []
                
                self.conversations[doc_id].extend([
                    {"role": "user", "content": query},
                    {"role": "assistant", "content": response_text}
                ])
            
            logger.info(f"Query processed successfully across documents {document_ids}")
            
            return {
                "response": response_text,
                "sources": sources
            }
            
        except Exception as e:
            logger.error(f"Error processing query across documents {document_ids}: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Error processing query: {str(e)}")

    def get_conversation_history(self, document_id: str) -> List[Dict]:
        """Get the conversation history for a specific document."""
        try:
            return self.conversations.get(document_id, [])
        except Exception as e:
            logger.error(f"Error retrieving conversation history for document {document_id}: {str(e)}")
            raise HTTPException(status_code=500, detail="Error retrieving conversation history")

    async def delete_document(self, document_id: str) -> None:
        """Delete a document from both the file system and vector database."""
        try:
            logger.info(f"Deleting document {document_id}")
            
            # Find and delete the file from the file system
            for file_path in self.documents_dir.glob(f"{document_id}_*"):
                if file_path.is_file():
                    file_path.unlink()
                    logger.info(f"Deleted file: {file_path}")
            
            # Delete document chunks from vector store
            self.vector_store._collection.delete(
                where={"doc_id": document_id}
            )
            
            # Remove conversation history
            if document_id in self.conversations:
                del self.conversations[document_id]
            
            logger.info(f"Document {document_id} deleted successfully")
            
        except Exception as e:
            logger.error(f"Error deleting document {document_id}: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Error deleting document: {str(e)}") 